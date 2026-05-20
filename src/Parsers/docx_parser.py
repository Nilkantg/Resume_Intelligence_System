from pathlib import Path
from typing import Dict, Any, Union
from docx import Document
from .base import BaseParser


class DOCXParser(BaseParser):
    """DOCX Parser"""

    def parse(self, file_path: Union[str, Path]) -> Dict[str, Any]:
        file_path = Path(file_path)
        if not file_path.exists():
            raise FileNotFoundError(f"File not found: {file_path}")

        doc = Document(file_path)
        
        full_text = []
        paragraphs = []
        
        for para in doc.paragraphs:
            if para.text.strip():
                paragraphs.append({
                    "text": para.text.strip(),
                    "style": para.style.name
                })
                full_text.append(para.text)

        # Extract tables
        tables = []
        for table in doc.tables:
            table_data = []
            for row in table.rows:
                row_data = [cell.text.strip() for cell in row.cells]
                table_data.append(row_data)
            if table_data:
                tables.append(table_data)

        result = {
            "file_name": file_path.name,
            "file_type": "docx",
            "text": "\n\n".join(full_text).strip(),
            "markdown": self._convert_to_markdown(full_text, tables),
            "paragraphs": paragraphs,
            "tables": tables,
            "num_paragraphs": len(paragraphs),
            "metadata": {"author": doc.core_properties.author}
        }
        
        return result

    def _convert_to_markdown(self, paragraphs: list, tables: list) -> str:
        """Simple markdown conversion"""
        md = []
        for p in paragraphs:
            md.append(p)
        md.append("\n\n")
        
        for i, table in enumerate(tables):
            md.append(f"**Table {i+1}**\n")
            for row in table:
                md.append("| " + " | ".join(row) + " |")
            md.append("\n")
        
        return "\n".join(md)

    def get_text(self, file_path: Union[str, Path]) -> str:
        return self.parse(file_path)["text"]

    def get_markdown(self, file_path: Union[str, Path]) -> str:
        return self.parse(file_path)["markdown"]
